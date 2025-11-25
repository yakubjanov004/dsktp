from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from typing import Dict, Any, List
from datetime import datetime

def _fmt_money(v) -> str:
    try:
        return f"{float(v):,.0f}".replace(",", " ")
    except Exception:
        return "0"

class AKTGenerator:
    """
    Shablonsiz (template-siz) .docx hujjat generatori.
    Sizning AKT maketingizdagi bo‘limlar (sarlavha, ma'lumotlar, materiallar jadvali,
    baholash, imzolar) to‘liq noldan yaratiladi.
    """
    def __init__(self):
        pass

    def generate_akt(self, data: Dict[str, Any], materials: List[Dict[str, Any]], output_path: str) -> bool:
        try:
            doc = Document()

            # --- Sahifa sozlamalari (margins) ---
            section = doc.sections[0]
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(1.5)

            # --- Default shrift ---
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)
            # Ruscha matnlar uchun (Word moslik)
            style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # --- Sarlavha ---
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("АКТ\nо начале эксплуатации оборудования и предоставления услуг сети")
            run.bold = True
            run.font.size = Pt(14)

            # --- Shahar va sana (hozirgi sana) ---
            row = doc.add_paragraph()
            row.alignment = WD_ALIGN_PARAGRAPH.LEFT
            row.add_run("г. Ташкент    ").bold = True
            row2 = doc.add_paragraph()
            row2.alignment = WD_ALIGN_PARAGRAPH.LEFT
            row2.add_run(datetime.now().strftime("«%d» %m.%Y г.")).bold = True

            doc.add_paragraph()  # bo'sh qator

            # --- Asosiy matn (kompaniya vakili, shartnoma, buyruq, tekshiruv) ---
            p = doc.add_paragraph()
            p.add_run("Представитель ООО «ALFA CONNECT» ").bold = True
            p.add_run(f"{data.get('technician_name', '—')} ")
            p.add_run("на основании договора № ").bold = True
            p.add_run(f"{data.get('contract_number', '—')} ")
            p.add_run("и служебного распоряжения № ").bold = True
            p.add_run(f"{data.get('service_order_number', '—')} ")
            p.add_run(", произвел тестирование и подключение услуг и передал ниже перечисленное оборудование, а представитель ")
            p.add_run(f"OOO «{data.get('organization_name', '___________________')}» ").bold = True
            p.add_run("проверил и принял предоставленные услуги и оборудование.")

            # --- Diagnostika / Ish tavsifi ---
            diag = data.get('diagnostics', '') or data.get('description_ish', '')
            if diag:
                doc.add_paragraph().add_run("Диагностика Абонентской линии:").bold = True
                doc.add_paragraph(diag)

            doc.add_paragraph()  # bo'sh qator

            # --- Xizmatlar mosligi ---
            p = doc.add_paragraph()
            p.add_run("Проведены необходимые проверки функционирования оборудования. Проверки показали, что предоставленные услуги: Интернет и установленное оборудование соответствуют указанным в договоре.")

            # --- Ekspluatatsiya boshlanishi va manzil ---
            start_text = doc.add_paragraph()
            start_text.add_run("На основании вышеизложенного, абонент начал эксплуатацию оборудования и услуг с ").bold = True
            start_text.add_run(datetime.now().strftime("«%d» %m.%Y г. "))
            start_text.add_run("по адресу: ").bold = True
            start_text.add_run(f"{data.get('address', '—')}")

            # --- Mijoz ma'lumotlari ---
            doc.add_paragraph().add_run("Абонент: ").bold = True
            doc.add_paragraph(f"{data.get('client_name', '—')}  |  {data.get('client_phone', '—')}")

            doc.add_paragraph().add_run("Тариф: ").bold = True
            doc.add_paragraph(f"{data.get('tariff_name', '—')}")

            # --- Materiallar jadvali ---
            doc.add_paragraph().add_run("Наименование оборудования и расходных материалов").bold = True
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Наименование"
            hdr[1].text = "Ед. изм"
            hdr[2].text = "Кол-во"
            hdr[3].text = "Цена"
            hdr[4].text = "Сумма"

            total_sum = 0.0
            if materials:
                for m in materials:
                    name = str(m.get('material_name', '—'))
                    unit = str(m.get('unit', 'шт'))
                    qty = float(m.get('quantity', 1) or 1)
                    price = float(m.get('price', 0) or 0)
                    total = float(m.get('total_price', qty * price) or 0)
                    total_sum += total

                    row = table.add_row().cells
                    row[0].text = name
                    row[1].text = unit
                    row[2].text = str(int(qty) if qty.is_integer() else qty)
                    row[3].text = _fmt_money(price)
                    row[4].text = _fmt_money(total)
            else:
                row = table.add_row().cells
                row[0].text = "Материалы не использованы"
                row[1].text = "—"
                row[2].text = "0"
                row[3].text = "0"
                row[4].text = "0"

            # Jami
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.add_run(f"Итого: {_fmt_money(total_sum)} сум").bold = True

            doc.add_paragraph()  # bo'sh qator

            # --- Baholash bo'limi ---
            p = doc.add_paragraph()
            p.add_run("Уважаемый Абонент! Просим Вас оценить работу представителя ООО «ALFA CONNECT»").bold = True
            
            # Client rating ko'rsatish
            client_rating = data.get('client_rating', 0)
            if client_rating > 0:
                rating_text = "★" * client_rating + "☆" * (5 - client_rating)
                doc.add_paragraph(f"Оценка клиента: {rating_text} ({client_rating}/5)")
            else:
                doc.add_paragraph("5    4    3    2    1    0")

            # --- Izoh (Kommentariya) ---
            doc.add_paragraph().add_run("Комментарий клиента:").bold = True
            client_comment = data.get('client_comment', '')
            if client_comment:
                doc.add_paragraph(f'"{client_comment}"')
            else:
                doc.add_paragraph("Комментарий не предоставлен")

            doc.add_paragraph()  

            sign = doc.add_paragraph()
            sign.add_run("Представитель ООО «ALFA CONNECT»: ").bold = True
            sign.add_run(f"{data.get('technician_name', '—')}")

            sign2 = doc.add_paragraph()
            sign2.add_run("Абонент (Ф.И.О): ").bold = True
            sign2.add_run(f"{data.get('client_name', '—')}")

            # Sana (imzo kuni)
            pdate = doc.add_paragraph()
            pdate.add_run("Дата: ").bold = True
            pdate.add_run(datetime.now().strftime("«%d» %m.%Y г."))

            # --- Saqlash ---
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"Error generating AKT: {e}")
            return False
