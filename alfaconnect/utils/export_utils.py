import csv
import io
import logging
from typing import List, Dict, Any
from datetime import datetime
import asyncio

# Set up logger
logger = logging.getLogger(__name__)
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from reportlab.lib.pagesizes import letter, A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportUtils:
    """Utility class for exporting data in various formats"""
    
    @staticmethod
    def _normalize_string(value: Any) -> str:
        """Convert any value to string safely."""
        try:
            if value is None:
                return ""
            return str(value)
        except Exception:
            return ""
    
    @staticmethod
    def generate_csv(data: List[Dict[str, Any]], filename: str = None) -> io.StringIO:
        """Generate CSV format from data"""
        if not data:
            return io.StringIO()
            
        output = io.StringIO()
        if data:
            fieldnames = data[0].keys()
            writer = csv.DictWriter(output, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(data)
        
        output.seek(0)
        return output
    
    @staticmethod
    def to_csv(data: List[Dict[str, Any]], headers: List[str] = None) -> io.StringIO:
        """Generate CSV format from data with optional custom headers."""
        if not data:
            return io.StringIO()
            
        output = io.StringIO()
        
        if data:
            # Use custom headers if provided, otherwise use data keys
            if headers:
                fieldnames = headers
                # Create new data with only the specified headers
                filtered_data = []
                for row in data:
                    filtered_row = {}
                    for header in headers:
                        # Try to find matching key in row (case-insensitive)
                        value = None
                        for key, val in row.items():
                            if str(key).lower() == str(header).lower():
                                value = val
                                break
                        filtered_row[header] = value if value is not None else ""
                    filtered_data.append(filtered_row)
                data = filtered_data
            else:
                fieldnames = data[0].keys()
            
            writer = csv.DictWriter(output, fieldnames=fieldnames, lineterminator='\n')
            # Normalize header names
            normalized_header_row = {h: ExportUtils._normalize_string(h) for h in fieldnames}
            writer.writerow(normalized_header_row)
            # Normalize each row value
            for row in data:
                normalized_row = {k: ExportUtils._normalize_string(row.get(k, "")) for k in fieldnames}
                writer.writerow(normalized_row)
        
        output.seek(0)
        return output
    
    @staticmethod
    def generate_excel(data: List[Dict[str, Any]], sheet_name: str = "Data", title: str = None) -> io.BytesIO:
        """Generate Excel format from data"""
        from datetime import datetime
        
        # Create a deep copy of the data to avoid modifying the original
        import copy
        data = copy.deepcopy(data)
        
        # Convert timezone-aware datetime objects to timezone-naive
        for row in data:
            for key, value in row.items():
                if isinstance(value, datetime) and value.tzinfo is not None:
                    row[key] = value.replace(tzinfo=None)
        
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name
        
        current_row = 1
        
        # Add title if provided
        if title:
            ws.cell(row=current_row, column=1, value=title)
            title_cell = ws.cell(row=current_row, column=1)
            title_cell.font = Font(size=16, bold=True)
            title_cell.alignment = Alignment(horizontal='center')
            title_cell.fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            current_row += 2  # Skip a row after title
        
        if data:
            # Add headers
            headers = [ExportUtils._normalize_string(h) for h in list(data[0].keys())]
            for col_num, header in enumerate(headers, 1):
                cell = ws.cell(row=current_row, column=col_num, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
                cell.alignment = Alignment(horizontal='center')
            
            current_row += 1
            
            # Add data rows
            for row_data in data:
                for col_num, value in enumerate(row_data.values(), 1):
                    ws.cell(row=current_row, column=col_num, value=ExportUtils._normalize_string(value))
                current_row += 1
            
            # Auto-adjust column widths
            for col_num in range(1, len(headers) + 1):
                max_length = 0
                
                header = str(headers[col_num - 1])
                header_length = len(header)
                if header_length > max_length:
                    max_length = header_length
                
                # Check data lengths in this column
                for row_num in range(1, current_row):
                    try:
                        cell_value = ws.cell(row=row_num, column=col_num).value
                        if cell_value is not None:
                            cell_str = str(cell_value)
                            if len(cell_str) > 1000:
                                cell_str = cell_str[:1000] + '...'
                            cell_length = len(cell_str)
                            if cell_length > max_length:
                                max_length = min(cell_length, 100)  
                    except Exception as e:
                        continue
                
                try:
                    from openpyxl.utils import get_column_letter
                    column_letter = get_column_letter(col_num)
                    adjusted_width = min(max(max_length + 2, 10), 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
                except Exception as e:
                    pass        
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        return output
    
    @staticmethod
    def generate_word(data: List[Dict[str, Any]], title: str = "Export Hisoboti") -> io.BytesIO:
        """Generate Word document from data"""
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(ExportUtils._normalize_string(title), level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation date
        date_paragraph = doc.add_paragraph(f"Yaratilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        doc.add_paragraph()  # Empty line
        
        if data:
            # Create table
            headers = [ExportUtils._normalize_string(h) for h in list(data[0].keys())]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add headers
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = str(header).replace('_', ' ').title()
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Add data rows
            for row_data in data:
                row_cells = table.add_row().cells
                for i, value in enumerate(row_data.values()):
                    normalized_value = ExportUtils._normalize_string(value)
                    row_cells[i].text = normalized_value if normalized_value is not None else ""
        else:
            doc.add_paragraph("Export uchun ma'lumotlar mavjud emas.")
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    @staticmethod
    def generate_pdf(data: List[Dict[str, Any]], title: str = "Export Hisoboti") -> io.BytesIO:
        """Generate PDF document from data with proper Unicode support"""
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.fonts import addMapping
        import os
        
        output = io.BytesIO()
        # Use landscape to give more horizontal space for many columns
        doc = SimpleDocTemplate(output, pagesize=landscape(A4))
        
        # Register a Unicode-capable TTF font for Cyrillic/Unicode text where possible
        try:
            candidates = []
            # Project-local font (recommended to ship this file)
            candidates.append(os.path.join(os.path.dirname(__file__), 'fonts', 'DejaVuSans.ttf'))
            # Common Windows fonts
            candidates += [
                r"C:\\Windows\\Fonts\\DejaVuSans.ttf",
                r"C:\\Windows\\Fonts\\NotoSans-Regular.ttf",
                r"C:\\Windows\\Fonts\\arial.ttf",
            ]
            # Common Linux fonts
            candidates += [
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
                "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
            ]
            font_name = 'Helvetica'
            for p in candidates:
                if os.path.exists(p):
                    pdfmetrics.registerFont(TTFont('AppUnicode', p))
                    font_name = 'AppUnicode'
                    break
        except:
            font_name = 'Helvetica'
        
        # Styles for better Unicode support
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # Center alignment
            fontName=font_name
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=9,
            leading=11,
            wordWrap='CJK'
        )
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=10,
            leading=12,
            wordWrap='CJK'
        )
        
        story = []
        
        # Add title
        story.append(Paragraph(ExportUtils._normalize_string(title), title_style))
        story.append(Paragraph(f"Yaratilgan sana: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", normal_style))
        story.append(Spacer(1, 20))
        
        if data:
            # Prepare table data with proper encoding
            headers = [ExportUtils._normalize_string(h) for h in list(data[0].keys())]
            table_data = [[Paragraph(str(h), header_style) for h in headers]]
            
            for row in data:
                row_values = []
                for value in row.values():
                    normalized = ExportUtils._normalize_string(value)
                    row_values.append(Paragraph(normalized if normalized is not None else "", normal_style))
                table_data.append(row_values)
            
            # Fit columns to available width
            available_width = doc.width
            num_cols = max(1, len(headers))
            # Minimum reasonable width per column
            min_w = 50
            base_w = max(min_w, available_width / num_cols)
            col_widths = [base_w] * num_cols

            # Create table
            table = Table(table_data, colWidths=col_widths, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ('TOPPADDING', (0, 0), (-1, -1), 3),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            
            story.append(table)
        else:
            story.append(Paragraph("Export uchun ma'lumotlar mavjud emas.", normal_style))
        
        doc.build(story)
        output.seek(0)
        return output
    
    @staticmethod
    def get_filename_with_timestamp(base_name: str, extension: str) -> str:
        """Generate filename with timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{base_name}_{timestamp}.{extension}"
    
    @staticmethod
    def format_data_for_export(data: List[Dict[str, Any]], export_type: str = "inventory") -> List[Dict[str, Any]]:
        """Format data specifically for export based on type"""
        if not data:
            return []
        
        formatted_data = []
        
        for item in data:
            if export_type == "inventory":
                formatted_item = {
                    "ID": item.get('id', ''),
                    "Nomi": item.get('name', ''),
                    "Miqdori": item.get('quantity', 0),
                    "Narxi": f"{item.get('price', 0):.2f}" if item.get('price') else "0.00",
                    "Umumiy Qiymat": f"{(item.get('quantity', 0) * item.get('price', 0)):.2f}" if item.get('price') else "0.00",
                    "Seriya Raqami": item.get('serial_number', ''),
                    "Tavsifi": item.get('description', ''),
                    "Yaratilgan": item.get('created_at', '').strftime('%Y-%m-%d %H:%M:%S') if item.get('created_at') else '',
                    "Yangilangan": item.get('updated_at', '').strftime('%Y-%m-%d %H:%M:%S') if item.get('updated_at') else ''
                }
            elif export_type == "statistics":
                formatted_item = {
                    "Ko'rsatkich": item.get('metric', ''),
                    "Qiymat": item.get('value', ''),
                    "Davr": item.get('period', ''),
                    "Sana": item.get('date', '').strftime('%Y-%m-%d') if item.get('date') else ''
                }
            else:
                formatted_item = item
            
            formatted_data.append(formatted_item)
        
        return formatted_data
    
    def generate_orders_export(self, orders_data: List[Dict[str, Any]], format_type: str, title: str) -> bytes:
        """Generate export file for orders data"""
        try:
            if format_type == "csv":
                csv_output = self.to_csv(orders_data)
                # Use BOM so Excel opens UTF-8 correctly
                return csv_output.getvalue().encode('utf-8-sig')
            elif format_type == "xlsx":
                excel_output = self.generate_excel(orders_data, "Orders", title)
                return excel_output.getvalue()
            elif format_type == "docx":
                word_output = self.generate_word(orders_data, title)
                return word_output.getvalue()
            elif format_type == "pdf":
                pdf_output = self.generate_pdf(orders_data, title)
                return pdf_output.getvalue()
            else:
                return None
        except Exception as e:
            print(f"Export generation error: {e}")
            return None
    
    def generate_statistics_export(self, stats_data: Dict[str, Any], format_type: str, title: str) -> bytes:
        """Generate export file for statistics data"""
        try:
            # Convert statistics to list format for export
            stats_list = []
            
            # Basic statistics
            stats_list.append({
                'Parametr': 'Jami arizalar',
                'Qiymat': stats_data.get('total_orders', 0)
            })
            stats_list.append({
                'Parametr': 'Kutilayotgan arizalar',
                'Qiymat': stats_data.get('pending_orders', 0)
            })
            stats_list.append({
                'Parametr': 'Jarayondagi arizalar',
                'Qiymat': stats_data.get('in_progress_orders', 0)
            })
            stats_list.append({
                'Parametr': 'Yakunlangan arizalar',
                'Qiymat': stats_data.get('completed_orders', 0)
            })
            stats_list.append({
                'Parametr': 'Bekor qilingan arizalar',
                'Qiymat': stats_data.get('cancelled_orders', 0)
            })
            stats_list.append({
                'Parametr': 'Tayinlangan arizalar',
                'Qiymat': stats_data.get('assigned_orders', 0)
            })
            stats_list.append({
                'Parametr': 'Tayinlanmagan arizalar',
                'Qiymat': stats_data.get('unassigned_orders', 0)
            })
            
            # Add region statistics
            region_stats = stats_data.get('region_stats', [])
            if region_stats:
                stats_list.append({'Parametr': '', 'Qiymat': ''})  # Empty row
                stats_list.append({'Parametr': 'HUDUDLAR BO\'YICHA', 'Qiymat': ''})
                for region in region_stats:
                    stats_list.append({
                        'Parametr': f"  {region['region']}",
                        'Qiymat': region['count']
                    })
            
            # Add tariff statistics
            tariff_stats = stats_data.get('tariff_stats', [])
            if tariff_stats:
                stats_list.append({'Parametr': '', 'Qiymat': ''})  # Empty row
                stats_list.append({'Parametr': 'TARIFLAR BO\'YICHA', 'Qiymat': ''})
                for tariff in tariff_stats:
                    stats_list.append({
                        'Parametr': f"  {tariff['tariff']}",
                        'Qiymat': tariff['count']
                    })
            
            # Generate export using the same methods as orders
            if format_type == "csv":
                csv_output = self.to_csv(stats_list)
                # Use BOM so Excel opens UTF-8 correctly
                return csv_output.getvalue().encode('utf-8-sig')
            elif format_type == "xlsx":
                excel_output = self.generate_excel(stats_list, "Statistics", title)
                return excel_output.getvalue()
            elif format_type == "docx":
                word_output = self.generate_word(stats_list, title)
                return word_output.getvalue()
            elif format_type == "pdf":
                pdf_output = self.generate_pdf(stats_list, title)
                return pdf_output.getvalue()
            else:
                return None
        except Exception as e:
            print(f"Statistics export generation error: {e}")
            return None