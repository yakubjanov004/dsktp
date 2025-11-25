from aiogram import Router, F
from aiogram.types import Message, CallbackQuery, BufferedInputFile
from aiogram.fsm.context import FSMContext
from filters.role_filter import RoleFilter
from keyboards.controllers_buttons import (
    get_controller_export_types_keyboard, 
    get_controller_export_formats_keyboard,
    get_controller_time_period_keyboard
)
from database.controller.export import (
    get_controller_orders_for_export,
    get_controller_statistics_for_export,
    get_controller_employees_for_export,
)
from utils.export_utils import ExportUtils
from utils.universal_error_logger import get_universal_logger, log_error
import logging
from datetime import datetime, timedelta

router = Router()
router.message.filter(RoleFilter(role="controller"))
logger = get_universal_logger("ControllerExport")

@router.message(F.text.in_(["📤 Export", "📤 Экспорт"]))
async def export_handler(message: Message, state: FSMContext):
    """Main export handler - shows export types"""
    try:
        await state.clear()
        keyboard = get_controller_export_types_keyboard()
        await message.answer(
            "📊 <b>Kontrollerlar uchun hisobotlar</b>\n\n"
            "Quyidagi hisobot turlaridan birini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
    except Exception as e:
        logger.error(f"Export handler error: {e}")
        await message.answer("❌ Xatolik yuz berdi. Iltimos, qaytadan urinib ko'ring.")

@router.callback_query(F.data == "controller_export_tech_requests")
async def export_tech_requests_handler(callback: CallbackQuery, state: FSMContext):
    """Handle tech requests export selection - show time period selection"""
    try:
        await state.update_data(export_type="tech_requests")
        from database.basic.language import get_user_language
        lang = await get_user_language(callback.from_user.id) or "uz"
        keyboard = get_controller_time_period_keyboard(lang)
        await callback.message.edit_text(
            "📋 <b>Texnik arizalar ro'yxati</b>\n\n"
            "Qaysi davr uchun export qilasiz?",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export tech requests handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)



@router.callback_query(F.data == "controller_export_statistics")
async def export_statistics_handler(callback: CallbackQuery, state: FSMContext):
    """Handle statistics export selection - show time period selection"""
    try:
        await state.update_data(export_type="statistics")
        from database.basic.language import get_user_language
        lang = await get_user_language(callback.from_user.id) or "uz"
        keyboard = get_controller_time_period_keyboard(lang)
        await callback.message.edit_text(
            "📊 <b>Statistika hisoboti</b>\n\n"
            "Qaysi davr uchun export qilasiz?",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export statistics handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

@router.callback_query(F.data == "controller_export_employees")
async def export_employees_handler(callback: CallbackQuery, state: FSMContext):
    """Handle employees export selection - show format selection directly"""
    try:
        await state.update_data(export_type="employees")
        from database.basic.language import get_user_language
        lang = await get_user_language(callback.from_user.id) or "uz"
        keyboard = get_controller_export_formats_keyboard(lang)
        await callback.message.edit_text(
            "👥 <b>Xodimlar ro'yxati</b>\n\n"
            "Barcha xodimlar (Controllerlar va Texniklar) export qilinadi.\n\n"
            "Export formatini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export employees handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

@router.callback_query(F.data.startswith("controller_time_"))
async def export_time_period_handler(callback: CallbackQuery, state: FSMContext):
    """Handle time period selection - show format selection"""
    try:
        time_period = callback.data.replace("controller_time_", "")  # today, week, month, total
        await state.update_data(time_period=time_period)
        
        from database.basic.language import get_user_language
        lang = await get_user_language(callback.from_user.id) or "uz"
        
        # Get period text
        period_texts = {
            "today": ("Bugungi hisobot", "Отчёт за сегодня"),
            "week": ("Haftalik hisobot (Dushanba - {today})", "Недельный отчёт (Понедельник - {today})"),
            "month": ("Oylik hisobot", "Месячный отчёт"),
            "total": ("Jami hisobot", "Общий отчёт")
        }
        
        export_type = (await state.get_data()).get("export_type", "tech_requests")
        
        # Calculate period text
        if time_period == "week":
            today = datetime.now().strftime("%d.%m.%Y")
            period_text = period_texts["week"][0].format(today=today) if lang == "uz" else period_texts["week"][1].format(today=today)
        else:
            period_text = period_texts[time_period][0] if lang == "uz" else period_texts[time_period][1]
        
        keyboard = get_controller_export_formats_keyboard(lang)
        
        title_text = {
            "tech_requests": ("Texnik arizalar ro'yxati", "Список технических заявок"),
            "statistics": ("Statistika hisoboti", "Статистический отчёт")
        }.get(export_type, ("Export", "Экспорт"))
        
        title = title_text[0] if lang == "uz" else title_text[1]
        
        await callback.message.edit_text(
            f"{get_emoji(export_type)} <b>{title}</b>\n\n"
            f"📅 Davr: <i>{period_text}</i>\n\n"
            f"Export formatini tanlang:",
            reply_markup=keyboard,
            parse_mode="HTML"
        )
        await callback.answer()
    except Exception as e:
        logger.error(f"Export time period handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

def get_emoji(export_type: str) -> str:
    """Get emoji for export type"""
    return {
        "tech_requests": "📋",
        "statistics": "📊",
        "employees": "👥"
    }.get(export_type, "📤")

@router.callback_query(F.data.startswith("controller_format_"))
async def export_format_handler(callback: CallbackQuery, state: FSMContext):
    """Handle export format selection and generate file"""
    try:
        format_type = callback.data.split("_")[-1]  # csv, xlsx, docx, pdf
        data = await state.get_data()
        export_type = data.get("export_type", "tech_requests")
        time_period = data.get("time_period", "total")  # today, week, month, total
        
        # Get data based on export type and time period
        if export_type == "tech_requests":
            orders_data = await get_controller_orders_for_export(time_period)
            title = "Texnik arizalar ro'yxati"
            filename_base = "texnik_arizalar"
            headers = [
                "ID", "Ariza raqami", "Mijoz ismi", "Telefon",
                "Manzil", "Ish tavsifi", "Holati",
                "Texnik", "Kontroller",
                "Yaratilgan sana", "Yangilangan sana",
                "Akt raqami"
            ]
            
            # Convert dict data to list format for export
            raw_data = [
                [
                    order.get("id", ""),
                    order.get("application_number", ""),
                    order.get("client_name", ""),
                    order.get("client_phone", ""),
                    order.get("address", ""),
                    order.get("description", ""),
                    order.get("status", ""),
                    order.get("technician_name", ""),  # Texnik
                    order.get("controller_name", ""),  # Kontroller
                    order.get("created_at", ""),
                    "",  # Yangilangan sana (bo'sh)
                    order.get("akt_number", "")  # Akt raqami
                ]
                for order in orders_data
            ]
            
        elif export_type == "connection_orders":
            raw_data = await get_controller_orders_for_export()
            title = "Ulanish arizalari ro'yxati"
            filename_base = "ulanish_arizalari"
            headers = [
                "ID", "Ariza raqami", "Mijoz ismi", "Telefon",
                "Manzil", "Tarif rejasi", "Holati",
                "Ulanish sanasi", "Yangilangan sana",
                "Akt raqami"
            ]
            
        elif export_type == "statistics":
            stats = await get_controller_statistics_for_export(time_period)

            if not stats or 'summary' not in stats:
                logger.error("Failed to get statistics for export or summary is missing.")
                await callback.message.answer(
                    "❌ Statistika ma'lumotlarini olishda xatolik yuz berdi.\n"
                    "Iltimos, keyinroq qayta urinib ko'ring."
                )
                await callback.answer()
                return
            raw_data = []
            title = "Statistika hisoboti"
            filename_base = "statistika"
            
            def add_section(title):
                nonlocal raw_data
                raw_data.append(["", ""])
                raw_data.append([f"🔹 {title.upper()}", ""])
                raw_data.append(["-" * 30, "-" * 30])
            
            def add_row(label, value, indent=0):
                nonlocal raw_data
                prefix = "  " * indent
                raw_data.append([f"{prefix}{label}", str(value) if value is not None else "0"])
            
            # 1. Asosiy statistika
            add_section("Umumiy statistika")
            add_row("📊 Jami texnik arizalar:", stats['summary']['total_requests'])
            add_row("🆕 Yangi arizalar:", stats['summary']['new_requests'])
            add_row("🔄 Jarayondagi arizalar:", stats['summary']['in_progress_requests'])
            add_row("✅ Yakunlangan arizalar:", stats['summary']['completed_requests'])
            add_row("📈 Yakunlangan arizalar foizi:", f"{stats['summary'].get('completion_rate', 0)}%")
            add_row("👥 Yagona mijozlar:", stats['summary']['unique_clients'])
            add_row("🔧 Muammo turlari:", stats['summary'].get('unique_tariffs', 0))
            
            # 2. Texniklar bo'yicha statistika
            if stats['by_technician']:
                add_section("Texniklar bo'yicha statistika")
                for i, technician in enumerate(stats['by_technician'], 1):
                    technician_name = f"👤 {i}. {technician['technician_name']}"
                    phone = technician['technician_phone'] or 'Tel. yo\'q'
                    add_row(technician_name, "", 0)
                    add_row("  📞 Telefon:", phone, 1)
                    add_row("  📊 Jami arizalar:", technician['total_orders'], 1)
                    add_row("  ✅ Yakunlangan:", technician['completed_orders'], 1)
                    add_row("  🔄 Jarayonda:", technician['in_progress_orders'], 1)
                    raw_data.append(["", ""])  # Empty row after each technician
            
            # 3. Oylik statistika
            if stats.get('monthly_trends'):
                add_section("Oylik statistika (6 oy)")
                for month_data in stats['monthly_trends']:
                    month = month_data['month']
                    add_row(f"🗓️ {month}:", "", 0)
                    add_row("  📊 Jami:", month_data['total_requests'], 1)
                    add_row("  🆕 Yangi:", month_data['new_requests'], 1)
                    add_row("  ✅ Yakunlangan:", month_data['completed_requests'], 1)
            
            # 4. Muammo turlari bo'yicha statistika (currently not implemented)
            # if stats['by_problem_type']:
            #     add_section("Muammo turlari bo'yicha statistika")
            #     for problem in stats['by_problem_type']:
            #         add_row(f"🔧 {problem['problem_type']}", "", 0)
            #         add_row("  📊 Arizalar soni:", problem['total_requests'], 1)
            #         add_row("  👥 Mijozlar soni:", problem['unique_clients'], 1)
            #         add_row("  ✅ Yakunlangan:", problem['completed_requests'], 1)
            
            # 5. So'nggi faollik
            if stats['recent_activity']:
                add_section("So'nggi faollik (30 kun)")
                for activity in stats['recent_activity']:
                    if activity['recent_orders'] > 0:
                        last_active = activity['last_order_date'].strftime('%Y-%m-%d') if activity['last_order_date'] else 'Noma\'lum'
                        add_row(
                            f"👤 {activity['technician_name']}",
                            f"📅 So'nggi: {last_active}",
                            0
                        )
                        add_row("  📊 Arizalar soni:", activity['recent_orders'], 1)
                
            headers = ["Ko'rsatkich", "Qiymat"]
            
        elif export_type == "employees":
            employees = await get_controller_employees_for_export()
            
            # Debug: log employees count
            logger.info(f"Employees from DB: {len(employees)}")
            
            title = "Xodimlar ro'yxati"
            filename_base = "xodimlar"
            headers = [
                "Ism-sharif", "Telefon", "Lavozim",
                "Qo'shilgan sana"
            ]
            
            seen_ids = set()
            unique_employees = []
            for emp in employees:
                emp_id = emp.get("id")
                if emp_id and emp_id not in seen_ids:
                    seen_ids.add(emp_id)
                    unique_employees.append(emp)
            
            logger.info(f"Unique employees after filtering: {len(unique_employees)}")
            
            raw_data = [
                [
                    emp.get("full_name", ""),
                    emp.get("phone", ""),
                    emp.get("role", ""),
                    emp.get("created_at", "").strftime('%Y-%m-%d') if emp.get("created_at") else ""
                ]
                for emp in unique_employees
            ]
            
            logger.info(f"Raw data rows: {len(raw_data)}")
            
        elif export_type == "reports":
            raw_data = await get_controller_reports_for_export()
            title = "Hisobotlar"
            filename_base = "hisobotlar"
            headers = [
                "Sarlavha", "Yaratuvchi", 
                "Holati", "Yaratilgan sana"
            ]
        
        # Generate file based on format
        try:
            if format_type == "xlsx":
                file = await generate_excel(raw_data, headers, title, filename_base)
            elif format_type == "csv":
                file = await generate_csv(raw_data, headers, title, filename_base)
            elif format_type == "docx":
                # For Word export, ensure data is in the correct format
                if export_type in ["employees", "reports"]:
                    # For these types, raw_data is already a list of dicts
                    file = await generate_word(raw_data, headers, title, filename_base)
                else:
                    # For other types, convert to list of dicts
                    dict_data = _rows_to_dicts(raw_data, headers)
                    file = await generate_word(dict_data, headers, title, filename_base)
            elif format_type == "pdf":
                file = await generate_pdf(raw_data, headers, title, filename_base)
            else:
                raise ValueError("Noto'g'ri format tanlandi")
        except Exception as e:
            logger.error(f"Error generating {format_type.upper()} file: {str(e)}", exc_info=True)
            raise ValueError(f"{format_type.upper()} faylini yaratishda xatolik: {str(e)}")
        
        # Format caption with time period (if applicable)
        from database.basic.language import get_user_language
        lang = await get_user_language(callback.from_user.id) or "uz"
        
        # Build caption
        if export_type == "employees":
            caption = f"📤 {title}\n" \
                     f"⏰ {datetime.now().strftime('%Y-%m-%d %H:%M')}\n" \
                     f"✅ Muvaffaqiyatli yuklab olindi!"
        else:
            period_texts = {
                "today": ("Bugun", "Сегодня"),
                "week": ("Hafta (Dushanba - hozirgi)", "Неделя (Понедельник - сейчас)"),
                "month": ("Oy", "Месяц"),
                "total": ("Jami", "Всего")
            }
            
            period_text = period_texts.get(time_period, ("Jami", "Всего"))[0] if lang == "uz" else period_texts.get(time_period, ("Jami", "Всего"))[1]
            
            caption = f"📤 {title}\n" \
                     f"📅 Davr: {period_text}\n" \
                     f"⏰ {datetime.now().strftime('%Y-%m-%d %H:%M')}\n" \
                     f"✅ Muvaffaqiyatli yuklab olindi!"
        
        # Send the file
        await callback.message.answer_document(
            document=file,
            caption=caption
        )
        
        # Remove the inline keyboard
        try:
            await callback.message.edit_reply_markup(reply_markup=None)
        except:
            pass
        
        await state.clear()
        
    except Exception as e:
        log_error(e, "Controller export format handler", callback.from_user.id)
        await callback.message.answer(
            f"❌ Xatolik yuz berdi: {str(e)}\n"
            "Iltimos, qaytadan urinib ko'ring yoki administratorga murojaat qiling."
        )
    
    await callback.answer()

@router.callback_query(F.data == "controller_export_back_types")
async def export_back_to_types_handler(callback: CallbackQuery, state: FSMContext):
    """Handle back - go to time period selection or initial export types"""
    try:
        data = await state.get_data()
        export_type = data.get("export_type")
        time_period = data.get("time_period")
        
        # If we're coming from format selection (time_period is set) and it's not employees,
        # go back to time period selection
        if time_period and export_type and export_type != "employees":
            # Remove time_period from state to allow re-selection
            await state.update_data(time_period=None)
            
            from database.basic.language import get_user_language
            lang = await get_user_language(callback.from_user.id) or "uz"
            keyboard = get_controller_time_period_keyboard(lang)
            
            title_text = {
                "tech_requests": ("Texnik arizalar ro'yxati", "Список технических заявок"),
                "statistics": ("Statistika hisoboti", "Статистический отчёт")
            }.get(export_type, ("Export", "Экспорт"))
            
            title = title_text[0] if lang == "uz" else title_text[1]
            emoji = get_emoji(export_type)
            
            edited = False
            try:
                await callback.message.edit_text(
                    f"{emoji} <b>{title}</b>\n\n"
                    f"Qaysi davr uchun export qilasiz?",
                    reply_markup=keyboard,
                    parse_mode="HTML"
                )
                edited = True
            except Exception as edit_error:
                if "message is not modified" in str(edit_error):
                    pass
                else:
                    raise edit_error
            
            if not edited:
                await callback.answer("✅", show_alert=False)
            else:
                await callback.answer()
        else:
            # Go back to initial export types screen
            # Clear state to reset the flow
            await state.clear()
            
            keyboard = get_controller_export_types_keyboard()
            edited = False
            try:
                await callback.message.edit_text(
                    "📊 <b>Kontrollerlar uchun hisobotlar</b>\n\n"
                    "Quyidagi hisobot turlaridan birini tanlang:",
                    reply_markup=keyboard,
                    parse_mode="HTML"
                )
                edited = True
            except Exception as edit_error:
                if "message is not modified" in str(edit_error):
                    pass
                else:
                    raise edit_error
            
            if not edited:
                await callback.answer("✅", show_alert=False)
            else:
                await callback.answer()
    except Exception as e:
        logger.error(f"Export back handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

@router.callback_query(F.data == "controller_export_end")
async def export_end_handler(callback: CallbackQuery, state: FSMContext):
    """Handle export end"""
    try:
        await callback.message.delete()
        await state.clear()
        await callback.answer()
    except Exception as e:
        logger.error(f"Export end handler error: {e}")
        await callback.answer("❌ Xatolik yuz berdi", show_alert=True)

def _rows_to_dicts(data: list, headers: list) -> list[dict]:
    """Convert rows to list of dicts with only the specified headers"""
    dict_data = []
    if not data:
        return dict_data

    # If rows are already dicts, just filter the keys
    if isinstance(data[0], dict):
        for row in data:
            row_dict = {}
            for header in headers:
                # Try direct key access first, then try with header mapping
                value = row.get(header) or row.get(_get_db_key_for_header(header), "")
                row_dict[header] = str(value) if value is not None else ""
            dict_data.append(row_dict)
        return dict_data

    # Handle list of lists/tuples
    for row in data:
        if not isinstance(row, (list, tuple, dict)):
            # Single value case
            row_dict = {headers[0]: str(row) if row is not None else ""}
            # Add empty values for remaining headers
            for header in headers[1:]:
                row_dict[header] = ""
        else:
            # List/tuple case
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    value = row[i]
                    row_dict[header] = str(value) if value is not None else ""
                else:
                    row_dict[header] = ""
        
        dict_data.append(row_dict)
    
    return dict_data


def _get_db_key_for_header(header: str) -> str:
    mapping = {
        # Common fields
        "ID": "id",
        "Holati": "status",
        "Yaratilgan sana": "created_at",
        "Yangilangan sana": "updated_at",
        
        # Client/technician/controller
        "Ariza raqami": "request_number",
        "Mijoz ismi": "client_name",
        "Telefon": "phone_number",
        "Mijoz abonent ID": "client_abonent_id",
        "Texnik": "assigned_technician",
        "Kontroller": "controller_name",
        
        # Address/work/tariff
        "Hudud": "region",
        "Manzil": "address",
        "Ish tavsifi": "description_ish",
        "Tarif rejasi": "plan_name",
        "Ulanish sanasi": "connection_date",
        
        # AKT documents
        "Akt raqami": "akt_number",
        "Akt fayl yo'li": "akt_file_path",
        "Akt yaratilgan": "akt_created_at",
        "Mijozga yuborilgan": "sent_to_client_at",
        "Akt reytingi": "akt_rating",
        "Akt izohi": "akt_comment",
        
        # Additional fields
        "Abonent ID": "abonent_id",
        "Media": "media",
        "Uzunlik": "longitude",
        "Kenglik": "latitude",
        "Muammo turi": "description",
        "Tavsif": "description",
        "Reyting": "rating",
        "Izohlar": "notes",
        "Texnik telefon": "technician_phone",
        "Kontroller telefon": "controller_phone"
    }
    return mapping.get(header, header.lower().replace(" ", "_")).replace(" ", "_")

async def generate_excel(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate Excel file from data"""
    logger.info(f"generate_excel called with {len(data)} rows of data for title: {title}")
    
    if not data:
        # Handle empty data
        dict_data = []
    else:
        # Data is in list format - convert to dict format
        dict_data = []
        for row in data:
            row_dict = {}
            for i, header in enumerate(headers):
                if i < len(row):
                    row_dict[header] = str(row[i]) if row[i] is not None else ""
                else:
                    row_dict[header] = ""
            dict_data.append(row_dict)
    
    logger.info(f"Converted to dict_data with {len(dict_data)} rows")
    
    # Use ExportUtils to generate Excel
    output = ExportUtils.generate_excel(dict_data, title[:30], title)
    
    return BufferedInputFile(
        file=output.getvalue(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    )

async def generate_csv(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate CSV file from data"""
    import io
    import csv
    
    # Convert data to list of dicts with only the specified headers
    dict_data = _rows_to_dicts(data, headers)
    
    output = io.StringIO()
    writer = csv.DictWriter(output, fieldnames=headers)
    
    # Write headers
    writer.writeheader()
    
    # Write data rows
    for row in dict_data:
        writer.writerow(row)
    
    return BufferedInputFile(
        file=output.getvalue(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.csv"
    )

async def generate_word(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate Word file from data"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    
    # Convert data to list of dicts with only the specified headers
    dict_data = _rows_to_dicts(data, headers)
    
    doc = Document()
    
    # Add title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Sana: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    date_run.italic = True
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    # Add data rows
    for row in dict_data:
        row_cells = table.add_row().cells
        for i, header in enumerate(headers):
            row_cells[i].text = str(row.get(header, ""))
    
    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return BufferedInputFile(
        file=output.read(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    )

async def generate_pdf(data: list, headers: list, title: str, filename: str) -> BufferedInputFile:
    """Generate PDF file from data"""
    # Convert data to list of dicts with only the specified headers
    dict_data = _rows_to_dicts(data, headers)
    
    # Use ExportUtils to generate PDF
    output = ExportUtils.generate_pdf(dict_data, title)
    
    return BufferedInputFile(
        file=output.getvalue(),
        filename=f"{filename}_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    )
